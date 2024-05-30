from docx import Document
from docx.shared import Inches

# Load the document
doc_path = "C:\Users\tppat\OneDrive\Documents\Scilab.docx"
doc = Document(doc_path)

# Function to change the page borders
def set_page_borders(doc):
    sections = doc.sections
    for section in sections:
        # Set the margin sizes to change the border
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

# Apply border changes
set_page_borders(doc)

# Function to remove comments and arrange the code properly
def clean_code_paragraphs(doc):
    for paragraph in doc.paragraphs:
        if "//" in paragraph.text:
            # Remove comments
            paragraph.text = paragraph.text.split("//")[0].strip()
        # Ensure proper spacing and new lines for readability
        if "clc;" in paragraph.text or "clear;" in paragraph.text:
            paragraph.add_run("\n")
        elif ";" in paragraph.text:
            paragraph.add_run("\n")

# Clean the code in the document
clean_code_paragraphs(doc)

# Save the updated document
output_path = "C:\Users\tppat\OneDrive\Documents\Scilab_final.docx"
doc.save(output_path)

output_path